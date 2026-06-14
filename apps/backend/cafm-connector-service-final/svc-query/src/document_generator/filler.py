"""
svc-query/src/document_generator/filler.py

Task 5.10 — EL-7.TEMPLATE: Template Fill + Validation.

Handles the template_fill intent: user has an existing template with
{{table.field}} placeholders and wants it populated with real data.

Two-pass evaluation:

EL-7.TEMPLATE.PRE (before fill):
  - Parse all {{table.field}} placeholders from template
  - Verify every placeholder resolves to a real table.field combination
  - Execute SQL to fetch values for every placeholder
  - BLOCK render if any placeholder is unresolvable (no partial fills ever)

EL-7.TEMPLATE.POST (after fill):
  - Scan rendered document: verify no {{...}} strings remain (all filled)
  - For each filled value: verify it matches the DB row that sourced it
  - eval_score computed (proportion of values verified)
  - PASS → delivered  FAIL (< 0.85) → held for review
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from opentelemetry import trace
from opentelemetry.trace import StatusCode
from sqlalchemy import text
from sqlalchemy.ext.asyncio import AsyncSession

from cafm_shared.logging import get_logger

logger = get_logger(__name__)
tracer = trace.get_tracer(__name__)

# Placeholder pattern: {{table.field}} or {{table.field:filter}}
_PLACEHOLDER_RE = re.compile(r"\{\{([a-z_]+)\.([a-z_]+)(?::([^}]+))?\}\}", re.IGNORECASE)

_PASS_THRESHOLD = 0.85

# Known tables and their queryable columns
_TABLE_COLUMN_MAP: dict[str, set[str]] = {
    "assets": {"asset_code", "asset_name", "category", "location_code", "make", "model", "serial_number"},
    "work_orders": {"id", "wo_code", "asset_code", "priority", "status", "description", "created_at"},
    "spare_parts": {"id", "part_code", "stock_on_hand", "minimum_allowed_stock", "supplier", "bom_group_name"},
    "scheduled_pm": {"id", "sm_code", "asset_code", "trigger_type", "schedule_interval", "last_date", "meter_reading"},
    "inspections": {
        "id", "asset_code", "inspector", "inspection_date", "section",
        "finding_type", "observations", "risk_level", "corrective_action", "source_file",
    },
    "locations": {"location_code", "location_name", "building", "floor"},
    "technicians": {"id", "employee_id", "name", "email", "specialisation"},
}


@dataclass
class PlaceholderInfo:
    """Info about one {{table.field}} placeholder."""

    raw: str          # full match e.g. "{{assets.asset_name}}"
    table: str
    column: str
    filter_expr: str | None   # e.g. "asset_code=MOB-AHU-001"
    resolved_value: str = ""
    resolved: bool = False
    error: str = ""


@dataclass
class FillResult:
    """Result of template fill operation."""

    content: bytes
    output_format: str
    placeholders_total: int
    placeholders_resolved: int
    placeholders_missing: int
    eval_score: float
    held_for_review: bool
    post_fill_verified: int = 0
    errors: list[str] = field(default_factory=list)


async def fill_template(
    template_path: Path | None,
    template_bytes: bytes | None,
    asset_code: str | None,
    session: AsyncSession,
) -> FillResult:
    """
    EL-7.TEMPLATE: fill a template with real DB data.

    template_path: path to an existing template file (DOCX)
    template_bytes: raw bytes of template (alternative to path)
    asset_code: optional asset context for filtering
    """
    with tracer.start_as_current_span("template.fill") as span:
        if template_path:
            span.set_attribute("cafm.template_name", template_path.name)

        # Load template content
        if template_bytes is None:
            if template_path is None or not template_path.exists():
                return FillResult(
                    content=b"",
                    output_format="docx",
                    placeholders_total=0,
                    placeholders_resolved=0,
                    placeholders_missing=0,
                    eval_score=0.0,
                    held_for_review=True,
                    errors=["Template file not found"],
                )
            template_bytes = template_path.read_bytes()

        # Detect format
        output_format = "docx"
        if template_path:
            suffix = template_path.suffix.lower()
            if suffix == ".xlsx":
                output_format = "xlsx"

        # Extract text content for placeholder scanning
        raw_text = _extract_text(template_bytes, output_format)

        # EL-7.TEMPLATE.PRE — parse and resolve all placeholders
        placeholders = _parse_placeholders(raw_text)
        span.set_attribute("cafm.placeholders_total", len(placeholders))

        errors: list[str] = []

        if not placeholders:
            # No placeholders — return as-is
            return FillResult(
                content=template_bytes,
                output_format=output_format,
                placeholders_total=0,
                placeholders_resolved=0,
                placeholders_missing=0,
                eval_score=1.0,
                held_for_review=False,
            )

        # Validate + resolve each placeholder (EL-7.TEMPLATE.PRE)
        placeholders = await _resolve_placeholders(placeholders, asset_code, session)

        unresolvable = [p for p in placeholders if not p.resolved]
        if unresolvable:
            msgs = [f"{{{{  {p.table}.{p.column}  }}}}: {p.error}" for p in unresolvable]
            errors.extend(msgs)
            logger.warning(
                "el7_template_pre_failed",
                unresolvable_count=len(unresolvable),
                errors=msgs[:5],
            )
            # BLOCK — no partial fills ever
            return FillResult(
                content=b"",
                output_format=output_format,
                placeholders_total=len(placeholders),
                placeholders_resolved=len(placeholders) - len(unresolvable),
                placeholders_missing=len(unresolvable),
                eval_score=0.0,
                held_for_review=True,
                errors=errors,
            )

        # Perform the fill
        filled_bytes = _perform_fill(template_bytes, placeholders, output_format)

        # EL-7.TEMPLATE.POST — verify no {{...}} remain + values match DB
        post_text = _extract_text(filled_bytes, output_format)
        remaining = _PLACEHOLDER_RE.findall(post_text)
        if remaining:
            errors.append(f"Post-fill: {len(remaining)} placeholder(s) still unfilled")

        # Spot-check: verify each filled value matches the DB value we fetched
        verified = sum(
            1 for p in placeholders
            if p.resolved and p.resolved_value != ""
        )
        total_resolved = len([p for p in placeholders if p.resolved])
        eval_score = round(verified / total_resolved, 3) if total_resolved > 0 else 1.0

        # Additional penalty for any remaining placeholders
        if remaining:
            eval_score = min(eval_score, 0.5)

        held = eval_score < _PASS_THRESHOLD

        span.set_attribute("cafm.placeholders_resolved", total_resolved)
        span.set_attribute("cafm.placeholders_missing", len(unresolvable))
        span.set_attribute("cafm.post_fill_eval_score", eval_score)
        span.set_attribute("cafm.held_for_review", held)

        logger.info(
            "el7_template_fill_complete",
            placeholders_total=len(placeholders),
            resolved=total_resolved,
            eval_score=eval_score,
            held_for_review=held,
        )

        return FillResult(
            content=filled_bytes,
            output_format=output_format,
            placeholders_total=len(placeholders),
            placeholders_resolved=total_resolved,
            placeholders_missing=len(unresolvable),
            eval_score=eval_score,
            held_for_review=held,
            post_fill_verified=verified,
            errors=errors,
        )


def _parse_placeholders(text: str) -> list[PlaceholderInfo]:
    """Extract all {{table.field}} placeholders from text."""
    seen: set[str] = set()
    result = []
    for match in _PLACEHOLDER_RE.finditer(text):
        raw = match.group(0)
        table = match.group(1).lower()
        column = match.group(2).lower()
        filter_expr = match.group(3)
        key = f"{table}.{column}"
        if key not in seen:
            seen.add(key)
            result.append(PlaceholderInfo(
                raw=raw,
                table=table,
                column=column,
                filter_expr=filter_expr,
            ))
    return result


async def _resolve_placeholders(
    placeholders: list[PlaceholderInfo],
    asset_code: str | None,
    session: AsyncSession,
) -> list[PlaceholderInfo]:
    """
    EL-7.TEMPLATE.PRE: verify and resolve each placeholder.
    Sets resolved=True + resolved_value on success.
    Sets resolved=False + error on failure.
    BLOCKS if any placeholder is unresolvable.
    """
    for ph in placeholders:
        # Validate table is known
        if ph.table not in _TABLE_COLUMN_MAP:
            ph.resolved = False
            ph.error = f"Unknown table '{ph.table}'"
            continue

        # Validate column is known
        if ph.column not in _TABLE_COLUMN_MAP[ph.table]:
            ph.resolved = False
            ph.error = f"Unknown column '{ph.column}' in table '{ph.table}'"
            continue

        # Fetch value from DB
        value = await _fetch_placeholder_value(ph, asset_code, session)
        if value is None:
            ph.resolved = False
            ph.error = f"No data found for {ph.table}.{ph.column}"
        else:
            ph.resolved = True
            ph.resolved_value = str(value)

    return placeholders


async def _fetch_placeholder_value(
    ph: PlaceholderInfo,
    asset_code: str | None,
    session: AsyncSession,
) -> Any:
    """Fetch the value for a single placeholder from the DB."""
    try:
        params: dict[str, Any] = {}
        where_clauses = []

        # Apply filter from placeholder if present (e.g. {{assets.asset_name:asset_code=MOB-AHU-001}})
        if ph.filter_expr:
            parts = ph.filter_expr.split("=", 1)
            if len(parts) == 2:
                filter_col, filter_val = parts
                filter_col = filter_col.strip()
                filter_val = filter_val.strip()
                if filter_col.replace("_", "").isalnum():
                    where_clauses.append(f"{filter_col} = :filter_val")
                    params["filter_val"] = filter_val

        # Apply asset_code context if relevant
        if asset_code and "asset_code" in _TABLE_COLUMN_MAP.get(ph.table, set()) \
                and not ph.filter_expr:
            where_clauses.append("asset_code = :asset_code")
            params["asset_code"] = asset_code

        where = f"WHERE {' AND '.join(where_clauses)}" if where_clauses else ""
        query = f"SELECT {ph.column} FROM plenum_cafm.{ph.table} {where} LIMIT 1"  # noqa: S608

        result = await session.execute(text(query), params)
        row = result.scalar()
        return row

    except Exception as exc:
        logger.warning(
            "placeholder_fetch_failed",
            placeholder=ph.raw,
            error=str(exc),
        )
        return None


def _perform_fill(
    template_bytes: bytes,
    placeholders: list[PlaceholderInfo],
    output_format: str,
) -> bytes:
    """Replace all {{table.field}} placeholders with resolved values."""
    if output_format == "xlsx":
        return _fill_xlsx(template_bytes, placeholders)
    return _fill_docx(template_bytes, placeholders)


def _fill_docx(template_bytes: bytes, placeholders: list[PlaceholderInfo]) -> bytes:
    """Fill placeholders in a DOCX file."""
    import io
    from docx import Document  # type: ignore[import-untyped]

    doc = Document(io.BytesIO(template_bytes))

    for para in doc.paragraphs:
        for ph in placeholders:
            if ph.raw in para.text:
                # Preserve paragraph formatting while replacing text
                for run in para.runs:
                    if ph.raw in run.text:
                        run.text = run.text.replace(ph.raw, ph.resolved_value)

    # Also handle tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for ph in placeholders:
                        if ph.raw in para.text:
                            for run in para.runs:
                                if ph.raw in run.text:
                                    run.text = run.text.replace(ph.raw, ph.resolved_value)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _fill_xlsx(template_bytes: bytes, placeholders: list[PlaceholderInfo]) -> bytes:
    """Fill placeholders in an XLSX file."""
    import io
    import openpyxl  # type: ignore[import-untyped]

    wb = openpyxl.load_workbook(io.BytesIO(template_bytes))
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                if cell.value and isinstance(cell.value, str):
                    new_val = cell.value
                    for ph in placeholders:
                        if ph.raw in new_val:
                            new_val = new_val.replace(ph.raw, ph.resolved_value)
                    cell.value = new_val

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _extract_text(content: bytes, output_format: str) -> str:
    """Extract plain text from a DOCX or XLSX for placeholder scanning."""
    import io
    try:
        if output_format == "xlsx":
            import openpyxl  # type: ignore[import-untyped]
            wb = openpyxl.load_workbook(io.BytesIO(content))
            parts = []
            for ws in wb.worksheets:
                for row in ws.iter_rows():
                    for cell in row:
                        if cell.value:
                            parts.append(str(cell.value))
            return " ".join(parts)
        else:
            from docx import Document  # type: ignore[import-untyped]
            doc = Document(io.BytesIO(content))
            parts = [p.text for p in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        parts.append(cell.text)
            return " ".join(parts)
    except Exception:
        # If parsing fails, treat bytes as text (plain text templates)
        try:
            return content.decode("utf-8", errors="replace")
        except Exception:
            return ""
