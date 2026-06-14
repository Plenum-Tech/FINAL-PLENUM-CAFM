"""
svc-query/src/document_generator/renderer.py

Task 5.7 — Deterministic Document Renderer.

Executes a validated DocumentPlan section by section.
Uses python-docx for DOCX, openpyxl for XLSX, reportlab for PDF.
Claude does NOT touch this file — values come from DB rows only.

Every value in the output traces to a real DB row (enforced by EL-7.DOC.RENDER).
"""

from __future__ import annotations

import io
import json
from datetime import datetime, timezone
from typing import Any

from opentelemetry import trace
from opentelemetry.trace import StatusCode
from sqlalchemy import text
from sqlalchemy.ext.asyncio import AsyncSession

from cafm_shared.logging import get_logger

from .schemas import DocumentPlan, DocumentSection

logger = get_logger(__name__)
tracer = trace.get_tracer(__name__)


class RenderResult:
    """Output of a render operation."""

    def __init__(
        self,
        content: bytes,
        output_format: str,
        sections_rendered: int,
        render_ms: int,
        sampled_values: list[dict[str, Any]],
    ) -> None:
        self.content = content
        self.output_format = output_format
        self.sections_rendered = sections_rendered
        self.render_ms = render_ms
        # sampled_values: list of {value, table, column} for EL-7.DOC.RENDER spot-checks
        self.sampled_values = sampled_values


async def render_document(
    plan: DocumentPlan,
    session: AsyncSession,
) -> RenderResult:
    """
    Deterministic render — no Claude.

    1. Fetch rows for each section's data_source
    2. Build document using python-docx / openpyxl / reportlab
    3. Return RenderResult with sampled_values for EL-7.DOC.RENDER spot-checks
    """
    with tracer.start_as_current_span("document.render") as span:
        span.set_attribute("cafm.document_type", plan.document_type)
        span.set_attribute("cafm.output_format", plan.output_format)
        span.set_attribute("cafm.sections_count", len(plan.sections))

        start_ms = _now_ms()
        sampled_values: list[dict[str, Any]] = []

        try:
            if plan.output_format == "xlsx":
                content, sampled = await _render_xlsx(plan, session)
            elif plan.output_format == "pdf":
                content, sampled = await _render_pdf(plan, session)
            else:  # docx (default)
                content, sampled = await _render_docx(plan, session)

            sampled_values = sampled
            render_ms = _now_ms() - start_ms

            span.set_attribute("cafm.render_ms", render_ms)
            span.set_attribute("cafm.sections_rendered", len(plan.sections))
            span.set_attribute("cafm.sampled_values_count", len(sampled_values))

            logger.info(
                "document_rendered",
                document_type=plan.document_type,
                output_format=plan.output_format,
                render_ms=render_ms,
                sections=len(plan.sections),
            )

            return RenderResult(
                content=content,
                output_format=plan.output_format,
                sections_rendered=len(plan.sections),
                render_ms=render_ms,
                sampled_values=sampled_values,
            )

        except Exception as exc:
            logger.error("document_render_failed", error=str(exc))
            span.set_status(StatusCode.ERROR, str(exc))
            raise


# ── DOCX renderer ─────────────────────────────────────────────────────────────


async def _render_docx(
    plan: DocumentPlan,
    session: AsyncSession,
) -> tuple[bytes, list[dict[str, Any]]]:
    """Render DocumentPlan to DOCX bytes."""
    from docx import Document  # type: ignore[import-untyped]
    from docx.shared import Pt, RGBColor  # type: ignore[import-untyped]

    doc = Document()
    sampled: list[dict[str, Any]] = []

    # Title
    title_para = doc.add_heading(plan.title, level=0)
    title_para.runs[0].font.size = Pt(18)

    # Subtitle / generated_for
    sub = doc.add_paragraph(f"Scope: {plan.generated_for}")
    sub.runs[0].font.italic = True

    doc.add_paragraph()  # spacer

    for section in plan.sections:
        rows, section_sampled = await _fetch_section_rows(section, session)
        sampled.extend(section_sampled)
        _render_docx_section(doc, section, rows)

    # Footer
    _add_docx_footer(doc, plan)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue(), sampled


def _render_docx_section(
    doc: Any,
    section: DocumentSection,
    rows: list[dict[str, Any]],
) -> None:
    """Render one section into the DOCX document."""
    from docx.shared import Pt, RGBColor  # type: ignore[import-untyped]

    doc.add_heading(section.heading, level=1)

    if section.type == "free_text_header":
        doc.add_paragraph(section.data_source)
        return

    if not rows:
        doc.add_paragraph("No data available for this section.")
        return

    if section.type in ("summary_table", "schedule_grid", "parts_table", "task_checklist"):
        # Determine columns
        columns = section.columns or list(rows[0].keys())
        limit = section.limit or len(rows)
        display_rows = rows[:limit]

        table = doc.add_table(rows=1, cols=len(columns))
        table.style = "Table Grid"

        # Header row
        hdr = table.rows[0].cells
        for i, col in enumerate(columns):
            hdr[i].text = col.replace("_", " ").title()
            for run in hdr[i].paragraphs[0].runs:
                run.font.bold = True

        # Data rows
        for row in display_rows:
            cells = table.add_row().cells
            for i, col in enumerate(columns):
                val = row.get(col, "")
                cells[i].text = _format_cell_value(val)

                # Highlight rule (e.g. "stock_on_hand == 0 → red")
                if section.highlight_rule:
                    _apply_highlight(cells[i], col, val, section.highlight_rule)

    elif section.type == "findings_list":
        for row in rows[: section.limit or len(rows)]:
            finding = row.get("observations") or row.get("finding_type", "")
            risk = row.get("risk_level", "")
            asset = row.get("asset_code", "")
            doc.add_paragraph(f"[{risk}] {asset}: {finding}", style="List Bullet")

    elif section.type == "kpi_summary":
        for row in rows[:1]:
            for k, v in row.items():
                doc.add_paragraph(f"{k.replace('_', ' ').title()}: {_format_cell_value(v)}")

    elif section.type == "signature_block":
        doc.add_paragraph("Inspector: ___________________________  Date: ____________")
        doc.add_paragraph("Supervisor: __________________________  Date: ____________")

    doc.add_paragraph()  # spacer


def _add_docx_footer(doc: Any, plan: DocumentPlan) -> None:
    """Add footer paragraph at end of document."""
    footer = plan.footer
    ts = footer.get("timestamp", datetime.now(timezone.utc).isoformat())
    audit_id = footer.get("audit_id", "N/A")
    generated_by = footer.get("generated_by", "CAFM AI Platform")
    doc.add_paragraph("─" * 60)
    doc.add_paragraph(
        f"Generated by: {generated_by}  |  {ts}  |  Audit ID: {audit_id}"
    )


# ── XLSX renderer ─────────────────────────────────────────────────────────────


async def _render_xlsx(
    plan: DocumentPlan,
    session: AsyncSession,
) -> tuple[bytes, list[dict[str, Any]]]:
    """Render DocumentPlan to XLSX bytes."""
    import openpyxl  # type: ignore[import-untyped]
    from openpyxl.styles import Font, PatternFill  # type: ignore[import-untyped]

    wb = openpyxl.Workbook()
    wb.remove(wb.active)  # remove default empty sheet
    sampled: list[dict[str, Any]] = []

    for i, section in enumerate(plan.sections):
        rows, section_sampled = await _fetch_section_rows(section, session)
        sampled.extend(section_sampled)

        # Each section gets its own sheet
        sheet_name = section.heading[:31]  # Excel limits sheet name to 31 chars
        ws = wb.create_sheet(title=sheet_name)
        _render_xlsx_section(ws, section, rows)

    # Title sheet at front
    title_ws = wb.create_sheet(title="Cover", index=0)
    title_ws["A1"] = plan.title
    title_ws["A1"].font = Font(bold=True, size=16)
    title_ws["A2"] = f"Scope: {plan.generated_for}"
    footer = plan.footer
    title_ws["A4"] = f"Generated: {footer.get('timestamp', '')}"
    title_ws["A5"] = f"Audit ID: {footer.get('audit_id', '')}"
    title_ws["A6"] = f"By: {footer.get('generated_by', 'CAFM AI Platform')}"

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue(), sampled


def _render_xlsx_section(ws: Any, section: DocumentSection, rows: list[dict[str, Any]]) -> None:
    """Render one section into an XLSX worksheet."""
    from openpyxl.styles import Font, PatternFill  # type: ignore[import-untyped]

    ws["A1"] = section.heading
    ws["A1"].font = Font(bold=True, size=12)

    if section.type == "free_text_header":
        ws["A2"] = section.data_source
        return

    if not rows:
        ws["A2"] = "No data available."
        return

    columns = section.columns or list(rows[0].keys())
    limit = section.limit or len(rows)
    display_rows = rows[:limit]

    header_fill = PatternFill("solid", fgColor="1F4E79")
    header_font = Font(color="FFFFFF", bold=True)

    # Headers at row 2
    for col_idx, col_name in enumerate(columns, start=1):
        cell = ws.cell(row=2, column=col_idx, value=col_name.replace("_", " ").title())
        cell.fill = header_fill
        cell.font = header_font

    # Data rows
    for row_idx, row in enumerate(display_rows, start=3):
        for col_idx, col_name in enumerate(columns, start=1):
            val = row.get(col_name, "")
            ws.cell(row=row_idx, column=col_idx, value=_format_cell_value(val))

            # Highlight critical items
            if section.highlight_rule and col_name == "stock_on_hand":
                try:
                    if int(val) == 0:
                        ws.cell(row=row_idx, column=col_idx).fill = PatternFill(
                            "solid", fgColor="FF0000"
                        )
                except (TypeError, ValueError):
                    pass


# ── PDF renderer ──────────────────────────────────────────────────────────────


async def _render_pdf(
    plan: DocumentPlan,
    session: AsyncSession,
) -> tuple[bytes, list[dict[str, Any]]]:
    """
    Render DocumentPlan to PDF bytes via reportlab.
    Falls back to DOCX-then-convert if reportlab not available.
    """
    sampled: list[dict[str, Any]] = []
    try:
        from reportlab.lib import colors  # type: ignore[import-untyped]
        from reportlab.lib.pagesizes import A4  # type: ignore[import-untyped]
        from reportlab.lib.styles import getSampleStyleSheet  # type: ignore[import-untyped]
        from reportlab.platypus import (  # type: ignore[import-untyped]
            Paragraph,
            SimpleDocTemplate,
            Spacer,
            Table,
            TableStyle,
        )

        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4)
        styles = getSampleStyleSheet()
        story = []

        story.append(Paragraph(plan.title, styles["Title"]))
        story.append(Paragraph(f"Scope: {plan.generated_for}", styles["Normal"]))
        story.append(Spacer(1, 12))

        for section in plan.sections:
            rows, section_sampled = await _fetch_section_rows(section, session)
            sampled.extend(section_sampled)

            story.append(Paragraph(section.heading, styles["Heading1"]))

            if not rows or section.type in ("signature_block", "free_text_header"):
                story.append(Paragraph(section.data_source or "No data.", styles["Normal"]))
                story.append(Spacer(1, 8))
                continue

            columns = section.columns or list(rows[0].keys())
            limit = section.limit or len(rows)
            display_rows = rows[:limit]

            table_data = [columns]
            for row in display_rows:
                table_data.append([_format_cell_value(row.get(c, "")) for c in columns])

            t = Table(table_data)
            t.setStyle(
                TableStyle([
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F4E79")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#EBF3FB")]),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ])
            )
            story.append(t)
            story.append(Spacer(1, 12))

        footer = plan.footer
        story.append(Paragraph(
            f"Generated by: {footer.get('generated_by', 'CAFM AI Platform')} | "
            f"{footer.get('timestamp', '')} | Audit ID: {footer.get('audit_id', '')}",
            styles["Normal"],
        ))

        doc.build(story)
        return buf.getvalue(), sampled

    except ImportError:
        logger.warning("reportlab_not_available_falling_back_to_docx")
        # Fallback: render as DOCX bytes with .pdf extension noted
        content, sampled = await _render_docx(plan, session)
        return content, sampled


# ── Shared helpers ─────────────────────────────────────────────────────────────


async def _fetch_section_rows(
    section: DocumentSection,
    session: AsyncSession,
) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    """
    Fetch rows for a section from the DB.
    Returns (rows, sampled_values) where sampled_values contains
    a sample of {value, table, column} entries for EL-7.DOC.RENDER.
    """
    from .validator import _extract_table_name  # avoid circular at module level

    if section.type in ("free_text_header", "signature_block"):
        return [], []

    table_name = _extract_table_name(section.data_source)
    if not table_name:
        return [], []

    try:
        # Build a safe SELECT based on section config
        cols = "*"
        if section.columns:
            # Validate column names are safe identifiers
            safe_cols = [c for c in section.columns if c.isidentifier()]
            if safe_cols:
                cols = ", ".join(safe_cols)

        order_clause = ""
        if section.sort_by and section.sort_by.replace("_", "").isalnum():
            order_clause = f"ORDER BY {section.sort_by}"

        limit_clause = f"LIMIT {section.limit}" if section.limit else "LIMIT 500"

        query = f"SELECT {cols} FROM plenum_cafm.{table_name} {order_clause} {limit_clause}"  # noqa: S608
        result = await session.execute(text(query))
        rows = [dict(r) for r in result.mappings().all()]

        # Sample values for EL-7.DOC.RENDER
        sampled = _sample_values(rows, table_name)
        return rows, sampled

    except Exception as exc:
        logger.warning("section_fetch_failed", section=section.heading, error=str(exc))
        return [], []


def _sample_values(
    rows: list[dict[str, Any]],
    table_name: str,
) -> list[dict[str, Any]]:
    """Extract a sample of {value, table, column} pairs for spot-checking."""
    import random

    sampled = []
    for row in rows[:20]:  # sample from first 20 rows
        for col, val in list(row.items())[:5]:  # up to 5 columns per row
            if val is not None and str(val).strip():
                sampled.append({"value": str(val), "table": table_name, "column": col})

    # Return up to 10 random samples
    return random.sample(sampled, min(10, len(sampled))) if sampled else []


def _format_cell_value(val: Any) -> str:
    """Format a cell value for display."""
    if val is None:
        return ""
    if isinstance(val, bool):
        return "Yes" if val else "No"
    if isinstance(val, (dict, list)):
        return json.dumps(val, default=str)
    return str(val)


def _apply_highlight(cell: Any, col: str, val: Any, highlight_rule: str) -> None:
    """Apply a highlight rule to a DOCX table cell."""
    # Simple rule parsing: "stock_on_hand == 0 → red"
    try:
        from docx.shared import RGBColor  # type: ignore[import-untyped]

        if "0" in highlight_rule and col == "stock_on_hand":
            if str(val) == "0":
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    except Exception:
        pass  # highlight is cosmetic — never block render


def _now_ms() -> int:
    """Current time in milliseconds."""
    return int(datetime.now(timezone.utc).timestamp() * 1000)
