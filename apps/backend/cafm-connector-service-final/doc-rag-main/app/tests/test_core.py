"""Basic unit tests. Run with: pytest app/tests/"""
from pathlib import Path

import pytest

from app.services.chunker import chunker
from app.services.document_classifier import document_classifier
from app.services.extraction_service import ExtractedDocument, ExtractedPage
from app.services.query_classifier import query_classifier
from app.utils.entity_extraction import extract_entities, extract_keys
from app.utils.text_normalization import normalize_key, normalize_text, token_overlap


# ---------- utils ----------
def test_normalize_text():
    assert normalize_text("  Hello   World  ") == "hello world"
    assert normalize_text("Café") == "cafe"


def test_normalize_key():
    assert normalize_key("AHU-17") == normalize_key("ahu 17") == normalize_key("AHU_17")


def test_token_overlap():
    assert token_overlap("the quick brown fox", "quick brown dog") > 0
    assert token_overlap("", "anything") == 0.0


def test_extract_entities():
    text = "Asset AHU-17 is covered by contract CTR-2025-009 under SLA-42."
    ents = extract_entities(text)
    assert "AHU-17" in ents.get("asset_code", [])
    keys = extract_keys(text)
    assert any("AHU" in k for k in keys)


# ---------- classifier ----------
def test_document_classifier_invoice():
    doc = ExtractedDocument(
        file_name="x.pdf", mime_type="application/pdf", num_pages=1,
        pages=[ExtractedPage(page_number=1, text=(
            "INVOICE NUMBER INV-123. Bill to Acme Corp. "
            "Line items below. Subtotal: $500. Amount due."
        ))],
    )
    doc_type, conf = document_classifier.classify(doc)
    assert doc_type == "invoice"
    assert conf > 0


def test_query_classifier_row_grounding():
    q = query_classifier.classify("What is the maintenance frequency for AHU-17?")
    assert q["query_type"] == "row_grounding"
    assert "AHU-17" in q["entity_keys"]


def test_query_classifier_comparison():
    q = query_classifier.classify("Compare SLA terms across vendor ABC and XYZ")
    assert q["query_type"] == "comparison"


# ---------- chunker ----------
def test_chunker_paragraph():
    doc = ExtractedDocument(
        file_name="x.txt", mime_type="text/plain", num_pages=1,
        pages=[ExtractedPage(
            page_number=1,
            text="Section: Termination. The agreement may be terminated with 30 days notice. " * 20,
        )],
    )
    chunks = chunker.chunk(doc)
    assert len(chunks) >= 1
    assert all(c.block_type == "paragraph" for c in chunks)


def test_chunker_table():
    doc = ExtractedDocument(
        file_name="x.txt", mime_type="text/plain", num_pages=1,
        pages=[ExtractedPage(
            page_number=1,
            text="Here is a preventive maintenance schedule.",
            tables=[[
                ["Asset", "Frequency", "Task"],
                ["AHU-17", "90 days", "Filter replacement"],
                ["AHU-18", "60 days", "Belt inspection"],
            ]],
        )],
    )
    chunks = chunker.chunk(doc)
    block_types = {c.block_type for c in chunks}
    assert "table_summary" in block_types
    assert "table_row" in block_types
    # Row chunks should contain the asset code
    row_chunks = [c for c in chunks if c.block_type == "table_row"]
    assert any("AHU-17" in c.text_content for c in row_chunks)


# ---------- multi-page extraction ----------
def test_docx_multipage_extraction(tmp_path):
    """Verify DOCX page breaks are honored and each page gets its own
    paragraphs + tables in the correct order."""
    from docx import Document as _Doc
    from app.services.extraction_service import extraction_service

    d = _Doc()
    d.add_paragraph("First page content with asset AHU-17.")
    tbl1 = d.add_table(rows=2, cols=2)
    tbl1.rows[0].cells[0].text = "Field"
    tbl1.rows[0].cells[1].text = "Value"
    tbl1.rows[1].cells[0].text = "Asset"
    tbl1.rows[1].cells[1].text = "AHU-17"

    d.add_page_break()
    d.add_paragraph("Second page talks about termination.")

    d.add_page_break()
    d.add_paragraph("Third page lists assets.")
    tbl2 = d.add_table(rows=2, cols=2)
    tbl2.rows[0].cells[0].text = "Asset"
    tbl2.rows[0].cells[1].text = "Frequency"
    tbl2.rows[1].cells[0].text = "AHU-18"
    tbl2.rows[1].cells[1].text = "60 days"

    fpath = tmp_path / "multi.docx"
    d.save(fpath)

    extracted = extraction_service.extract(fpath)

    assert extracted.num_pages == 3
    assert "AHU-17" in extracted.pages[0].text
    assert len(extracted.pages[0].tables) == 1
    assert "termination" in extracted.pages[1].text.lower()
    assert len(extracted.pages[1].tables) == 0
    assert "assets" in extracted.pages[2].text.lower()
    assert len(extracted.pages[2].tables) == 1
    # Page ordering is preserved
    assert [p.page_number for p in extracted.pages] == [1, 2, 3]


def test_vision_extraction_injects_tables_as_chunks(tmp_path, monkeypatch):
    """When the vision service finds a table inside an image, that table
    must become real `table_row` chunks in the final chunk list.

    We monkey-patch VisionService.extract_from_image with a fake that
    returns a structured table, run a DOCX with an embedded image
    through the pipeline, and verify table_row chunks appear.
    """
    from docx import Document as _Doc
    from app.services import vision_service as vm
    from app.services.vision_service import VisionResult
    from app.services.extraction_service import extraction_service
    from app.services.chunker import chunker

    # --- Build a DOCX with a real embedded image ---
    from PIL import Image as _PIL
    img_file = tmp_path / "tiny.png"
    _PIL.new("RGB", (200, 200), color=(200, 200, 200)).save(img_file, "PNG")

    d = _Doc()
    d.add_paragraph("Before the image.")
    d.add_picture(str(img_file))
    d.add_paragraph("After the image.")
    doc_path = tmp_path / "with_image.docx"
    d.save(doc_path)

    # --- Fake vision service: return a structured 2-row table ---
    def fake_extract(image_bytes, image_format="png"):
        return VisionResult(
            description="Test table inside image.",
            tables=[[
                ["Asset", "Frequency"],
                ["AHU-17", "90 days"],
                ["AHU-18", "60 days"],
            ]],
            extra_text="",
            entities={},
        )

    monkeypatch.setattr(vm.vision_service, "mock", False)
    monkeypatch.setattr(vm.vision_service, "_client", object())
    monkeypatch.setattr(vm.vision_service, "extract_from_image", fake_extract)
    # Also bypass the "too small" guard for the 1x1 test image
    monkeypatch.setattr(vm.vision_service, "should_skip", lambda w, h: False)

    extracted = extraction_service.extract(doc_path, document_id="test-doc")

    # One page, one image, one injected table
    assert extracted.num_pages == 1
    assert extracted.num_images == 1
    assert len(extracted.pages[0].tables) == 1
    assert extracted.pages[0].tables[0][0] == ["Asset", "Frequency"]

    # Chunker should turn the injected table into table_row chunks
    chunks = chunker.chunk(extracted)
    types = {c.block_type for c in chunks}
    assert "table_row" in types
    row_chunks = [c for c in chunks if c.block_type == "table_row"]
    assert len(row_chunks) == 2
    assert any("AHU-17" in c.text_content for c in row_chunks)
    assert any("AHU-18" in c.text_content for c in row_chunks)
    # Image chunk should record the vision method
    img_chunks = [c for c in chunks if c.block_type == "image"]
    assert len(img_chunks) == 1
    assert img_chunks[0].meta.get("extraction_method") == "vision"


# ---------- vision response parser leniency ----------
def test_vision_parser_canonical_shape():
    """Canonical {headers, rows} shape with image_type='table'."""
    from app.services.vision_service import vision_service
    raw = (
        '{"image_type":"table","description":"x","tables":[{"title":"t",'
        '"headers":["A","B"],"rows":[["1","2"],["3","4"]]}],'
        '"extra_text":"","entities":{}}'
    )
    r = vision_service._parse_response(raw)
    assert r.image_type == "table"
    assert len(r.tables) == 1
    assert r.tables[0] == [["A", "B"], ["1", "2"], ["3", "4"]]


def test_vision_parser_alt_key_names():
    """Model sometimes emits 'columns'/'data' instead of 'headers'/'rows'."""
    from app.services.vision_service import vision_service
    raw = (
        '{"image_type":"table","description":"x","tables":[{'
        '"columns":["A","B"],"data":[["1","2"],["3","4"]]}],'
        '"extra_text":"","entities":{}}'
    )
    r = vision_service._parse_response(raw)
    assert len(r.tables) == 1
    assert r.tables[0][0] == ["A", "B"]
    assert len(r.tables[0]) == 3  # headers + 2 data rows


def test_vision_parser_row_as_dict():
    """Row-as-object {col: val} shape should be re-keyed by headers."""
    from app.services.vision_service import vision_service
    raw = (
        '{"image_type":"table","description":"x","tables":[{'
        '"headers":["Asset","Frequency"],'
        '"rows":[{"Asset":"AHU-17","Frequency":"90 days"},'
        '{"Asset":"AHU-18","Frequency":"60 days"}]}],'
        '"extra_text":"","entities":{}}'
    )
    r = vision_service._parse_response(raw)
    assert len(r.tables) == 1
    assert r.tables[0] == [
        ["Asset", "Frequency"],
        ["AHU-17", "90 days"],
        ["AHU-18", "60 days"],
    ]


def test_vision_parser_single_table_as_dict():
    """Non-spec: `tables` is a single dict instead of a list of dicts."""
    from app.services.vision_service import vision_service
    raw = (
        '{"image_type":"table","description":"x",'
        '"tables":{"headers":["A","B"],"rows":[["1","2"]]},'
        '"extra_text":"","entities":{}}'
    )
    r = vision_service._parse_response(raw)
    assert len(r.tables) == 1
    assert r.tables[0] == [["A", "B"], ["1", "2"]]


def test_vision_parser_flat_list_of_lists():
    """Model sometimes emits a flat [[headers],[row]] with no wrapper dict."""
    from app.services.vision_service import vision_service
    raw = (
        '{"image_type":"table","description":"x",'
        '"tables":[[["A","B"],["1","2"],["3","4"]]],'
        '"extra_text":"","entities":{}}'
    )
    r = vision_service._parse_response(raw)
    assert len(r.tables) == 1
    # Flat shape is passed through as-is
    assert r.tables[0][0] == ["A", "B"]


def test_vision_parser_table_classified_but_empty_logs_error(caplog):
    """When image_type='table' but tables is empty, the parser logs
    an ERROR so the failure is visible in the operator logs."""
    import logging
    from app.services.vision_service import vision_service
    raw = '{"image_type":"table","description":"a table","tables":[],"extra_text":"","entities":{}}'
    with caplog.at_level(logging.ERROR):
        r = vision_service._parse_response(raw)
    assert r.image_type == "table"
    assert r.tables == []
    # raw_response is always preserved for later debugging
    assert r.raw_response == raw
