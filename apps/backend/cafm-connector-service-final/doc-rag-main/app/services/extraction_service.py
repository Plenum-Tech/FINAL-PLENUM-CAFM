"""Document extraction service.

Handles PDF (via Claude Vision), DOCX, and TXT.
Returns a normalized intermediate JSON of pages, paragraphs, tables, and images
— the input to the classifier and chunker.

PDF extraction uses Claude's native document understanding (base64 inline).
DOCX extraction uses python-docx + vision/OCR for embedded images.
"""
from __future__ import annotations

import base64
import json
import threading
import time
import uuid
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from app.core.config import settings
from app.core.logger import logger
from app.services.ocr_service import ocr_service
from app.services.vision_service import vision_service
from tenacity import retry, retry_if_exception, stop_after_attempt, wait_exponential

_PDF_EXTRACTION_PROMPT = """\
Extract the complete content from this PDF document. Return a JSON object with \
this exact structure — no prose, no markdown fences, just the JSON:

{
  "pages": [
    {
      "page_number": 1,
      "text": "<all non-table readable text on this page: headings, paragraphs, \
bullet points, captions, footnotes — preserve line breaks with \\n>",
      "tables": [
        [
          ["col1_header", "col2_header"],
          ["row1_val1",   "row1_val2"],
          ["row2_val1",   "row2_val2"]
        ]
      ]
    }
  ]
}

Rules:
- Include every page in document order.
- text: all non-table text. Do NOT include table content here.
- tables: each table as a list of rows; first row = column headers; \
preserve numbers and units exactly; repeat merged-cell values per row; \
return [] if the page has no tables.
- Return ONLY the JSON object.
"""

# Cap simultaneous Anthropic PDF extractions per process to reduce burst overload.
_PDF_EXTRACTION_SEMAPHORE = threading.BoundedSemaphore(3)


def _is_retryable_anthropic_error(exc: BaseException) -> bool:
    """Retry only on provider overload / throttling conditions."""
    status_code = getattr(exc, "status_code", None)
    if status_code in (429, 529):
        return True
    payload = str(getattr(exc, "body", "") or str(exc)).lower()
    return "overloaded_error" in payload or "rate_limit_error" in payload


@dataclass
class ExtractedImage:
    """An image extracted from a document page.

    - `ocr_text` is the text OCR'd from the image (may be empty).
    - `saved_path` is where the image bytes were written to disk so the
      UI can display it and so it can be referenced from chunks.
    - `url_path` is the public URL (relative to API base) at which the
      image is served by FastAPI's StaticFiles mount.
    """
    page_number: int
    image_index: int
    ocr_text: str = ""
    width: int = 0
    height: int = 0
    format: str = ""
    saved_path: str | None = None
    url_path: str | None = None
    bbox: tuple[float, float, float, float] | None = None
    source: str = "embedded"  # embedded | rendered_page


@dataclass
class ExtractedPage:
    page_number: int
    text: str
    tables: list[list[list[str]]] = field(default_factory=list)  # list of rows-of-cells
    images: list[ExtractedImage] = field(default_factory=list)


@dataclass
class ExtractedDocument:
    file_name: str
    mime_type: str
    num_pages: int
    pages: list[ExtractedPage] = field(default_factory=list)

    @property
    def full_text(self) -> str:
        return "\n\n".join(p.text for p in self.pages if p.text)

    @property
    def num_images(self) -> int:
        return sum(len(p.images) for p in self.pages)

    def to_dict(self) -> dict[str, Any]:
        return {
            "file_name": self.file_name,
            "mime_type": self.mime_type,
            "num_pages": self.num_pages,
            "num_images": self.num_images,
            "pages": [
                {
                    "page_number": p.page_number,
                    "text": p.text,
                    "tables": p.tables,
                    "images": [
                        {
                            "image_index": img.image_index,
                            "ocr_text": img.ocr_text,
                            "width": img.width,
                            "height": img.height,
                            "format": img.format,
                            "url_path": img.url_path,
                            "source": img.source,
                        }
                        for img in p.images
                    ],
                }
                for p in self.pages
            ],
        }


class ExtractionService:
    """PDF + DOCX extraction with table awareness."""

    IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp", ".tif", ".tiff", ".gif"}
    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", *IMAGE_EXTENSIONS}

    def extract(
        self,
        file_path: str | Path,
        document_id: str | None = None,
    ) -> ExtractedDocument:
        """Extract a document.

        If `document_id` is provided, embedded images are saved under
        `<upload_dir>/images/<document_id>/` and each ExtractedImage gets
        `saved_path` and `url_path` populated. If `document_id` is None,
        a temporary UUID is generated — images are still saved.
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {path}")

        ext = path.suffix.lower()
        doc_id = document_id or f"tmp-{uuid.uuid4().hex[:8]}"
        image_dir = self._image_dir(doc_id)
        logger.info(
            "Extracting document | file={} | ext={} | doc_id={} | ocr={}",
            path.name, ext, doc_id, ocr_service.available,
        )

        try:
            if ext == ".pdf":
                doc = self._extract_pdf(path, doc_id, image_dir)
            elif ext == ".docx":
                doc = self._extract_docx(path, doc_id, image_dir)
            elif ext == ".txt":
                doc = self._extract_txt(path)
            elif ext in self.IMAGE_EXTENSIONS:
                doc = self._extract_image(path, doc_id, image_dir)
            else:
                raise ValueError(
                    f"Unsupported file type: {ext}. "
                    f"Supported: {sorted(self.SUPPORTED_EXTENSIONS)}"
                )
        except Exception as e:
            logger.exception("Extraction failed | file={} | err={}", path.name, e)
            raise

        logger.info(
            "Extraction complete | file={} | pages={} | chars={} | images={}",
            path.name, doc.num_pages, len(doc.full_text), doc.num_images,
        )
        return doc

    # ---------- image helpers ----------
    def _image_dir(self, document_id: str) -> Path:
        base = Path(settings.upload_dir) / "images" / document_id
        base.mkdir(parents=True, exist_ok=True)
        return base

    def _save_image_bytes(
        self,
        data: bytes,
        image_dir: Path,
        document_id: str,
        page_number: int,
        image_index: int,
        fmt: str = "png",
    ) -> tuple[str, str]:
        """Write image bytes to disk. Returns (absolute_saved_path, url_path)."""
        filename = f"page{page_number}_img{image_index}.{fmt.lower()}"
        abs_path = image_dir / filename
        try:
            abs_path.write_bytes(data)
        except Exception as e:
            logger.warning("Failed to save image {}: {}", abs_path, e)
            return "", ""
        # URL served by FastAPI StaticFiles at /images/<doc_id>/<file>
        url_path = f"/images/{document_id}/{filename}"
        return str(abs_path), url_path

    def _process_image_bytes(
        self,
        data: bytes,
        fmt: str,
        page_number: int,
        image_index: int,
        doc_id: str,
        image_dir: Path,
        source: str = "embedded",
    ) -> tuple[ExtractedImage | None, list[list[list[str]]]]:
        """Save + analyze an image.

        Order of preference:
          1. OpenAI vision (if OPENAI_API_KEY is set) — returns a
             description + structured tables + entities.
          2. Tesseract OCR — returns raw text only.
          3. Neither — the image is saved with just metadata so the UI
             can still show it.

        Returns:
          - ExtractedImage (or None if save failed)
          - list of tables extracted from this image. Callers should
            append these to their page.tables so the chunker produces
            real table_row chunks for them.
        """
        saved_abs, url = self._save_image_bytes(
            data, image_dir, doc_id, page_number, image_index, fmt=fmt,
        )
        if not saved_abs:
            return None, []

        width = height = 0
        try:
            from io import BytesIO
            from PIL import Image
            pil = Image.open(BytesIO(data))
            width, height = pil.size
        except Exception as e:
            logger.warning("Failed to read image dimensions: {}", e)

        # Skip tiny decorative images (icons, bullets) — don't waste an API
        # call on them, but still keep them saved/indexable by metadata.
        too_small = vision_service.should_skip(width, height)

        ocr_text = ""
        description_text = ""
        injected_tables: list[list[list[str]]] = []
        entities: dict[str, list[str]] = {}
        method = "none"
        vision_image_type = ""
        vision_raw = ""

        if vision_service.available and not too_small:
            vr = vision_service.extract_from_image(data, image_format=fmt)
            vision_image_type = vr.image_type
            vision_raw = vr.raw_response
            if vr.description or vr.extra_text or vr.tables:
                description_text = vr.combined_text
                injected_tables = vr.tables
                entities = vr.entities
                method = "vision"
                logger.info(
                    "Vision extracted | page={} | img={} | type={} | "
                    "tables={} | desc_chars={}",
                    page_number, image_index, vr.image_type or "?",
                    len(vr.tables), len(description_text),
                )
            else:
                logger.warning(
                    "Vision returned empty result | page={} | img={} | "
                    "raw={}", page_number, image_index, vr.raw_response[:300],
                )

        if not description_text and ocr_service.available and not too_small:
            try:
                from io import BytesIO
                from PIL import Image
                pil = Image.open(BytesIO(data))
                ocr_text = ocr_service.ocr_pil(pil)
                if ocr_text:
                    method = "ocr"
                    logger.info(
                        "OCR extracted | page={} | img={} | chars={}",
                        page_number, image_index, len(ocr_text),
                    )
            except Exception as e:
                logger.warning("OCR fallback failed: {}", e)

        # Combine what we got: prefer vision description, then OCR text
        combined_text = description_text or ocr_text

        image = ExtractedImage(
            page_number=page_number,
            image_index=image_index,
            ocr_text=combined_text,
            width=width,
            height=height,
            format=fmt.upper(),
            saved_path=saved_abs,
            url_path=url,
            source=source,
        )
        # Stash the method on the image via side-channel attributes so the
        # chunker can persist them into the chunk's metadata JSON. These
        # become visible via GET /documents/{id}/chunks and via the new
        # /documents/{id}/pages/{n}/debug endpoint, so you can see exactly
        # what vision returned for any page that retrieved poorly.
        image.extraction_method = method  # type: ignore[attr-defined]
        image.entities = entities          # type: ignore[attr-defined]
        image.vision_image_type = vision_image_type  # type: ignore[attr-defined]
        image.vision_raw = vision_raw[:2000]         # type: ignore[attr-defined]
        image.injected_table_count = len(injected_tables)  # type: ignore[attr-defined]
        return image, injected_tables

    # ---------- PDF (Claude) ----------
    def _extract_pdf(
        self, path: Path, doc_id: str, image_dir: Path
    ) -> ExtractedDocument:
        return self._extract_pdf_claude(path)

    def _extract_pdf_claude(self, path: Path) -> ExtractedDocument:
        """Send the entire PDF to Claude and parse the structured JSON response.

        Claude reads text, tables, and images natively in one call — no
        separate pdfplumber, pypdf, or OCR pass needed.
        """
        if not settings.anthropic_api_key:
            raise RuntimeError(
                "ANTHROPIC_API_KEY is not set. Claude-based PDF extraction requires "
                "this variable to be configured."
            )

        try:
            import anthropic
        except ImportError:
            raise RuntimeError(
                "The 'anthropic' package is not installed. "
                "Run: pip install anthropic"
            )

        pdf_bytes = path.read_bytes()
        b64_data = base64.standard_b64encode(pdf_bytes).decode("ascii")
        size_mb = len(pdf_bytes) / (1024 * 1024)

        logger.info(
            "Claude PDF extraction starting | file={} | size_mb={:.2f} | model={}",
            path.name, size_mb, settings.claude_pdf_model,
        )

        client = anthropic.Anthropic(api_key=settings.anthropic_api_key)
        t0 = time.time()

        messages = [{
            "role": "user",
            "content": [
                {
                    "type": "document",
                    "source": {
                        "type": "base64",
                        "media_type": "application/pdf",
                        "data": b64_data,
                    },
                    "cache_control": {"type": "ephemeral"},
                },
                {
                    "type": "text",
                    "text": _PDF_EXTRACTION_PROMPT,
                },
            ],
        }]

        # Use streaming — required for long-running extraction calls.
        # Wrapped with targeted retry + per-process concurrency guard.
        final = self._call_claude_stream_with_retry(client=client, messages=messages)

        elapsed_ms = int((time.time() - t0) * 1000)
        raw = final.content[0].text if final.content else ""

        logger.info(
            "Claude PDF extraction done | file={} | ms={} | "
            "input_tokens={} | output_tokens={} | raw_chars={}",
            path.name, elapsed_ms,
            final.usage.input_tokens, final.usage.output_tokens, len(raw),
        )

        # Strip accidental markdown fences before parsing
        text_to_parse = raw.strip()
        if text_to_parse.startswith("```"):
            text_to_parse = text_to_parse.split("\n", 1)[-1]
            text_to_parse = text_to_parse.rsplit("```", 1)[0]

        try:
            data = json.loads(text_to_parse)
        except json.JSONDecodeError as e:
            logger.error(
                "Claude PDF response was not valid JSON | file={} | err={} | raw={}",
                path.name, e, raw[:500],
            )
            raise RuntimeError(
                f"Claude returned non-JSON response for {path.name}: {e}"
            ) from e

        raw_pages = data.get("pages") or []
        if not raw_pages:
            logger.warning("Claude returned no pages for {}", path.name)

        pages: list[ExtractedPage] = []
        empty_pages: list[int] = []
        for p in raw_pages:
            page_num = int(p.get("page_number") or len(pages) + 1)
            text = str(p.get("text") or "").strip()

            # Normalise tables — each table is [[row], [row], ...]
            raw_tables = p.get("tables") or []
            tables: list[list[list[str]]] = []
            for t in raw_tables:
                if not isinstance(t, list) or len(t) < 2:
                    continue
                cleaned = [
                    [str(cell).strip() for cell in row]
                    for row in t
                    if isinstance(row, list) and any(str(c).strip() for c in row)
                ]
                if cleaned:
                    tables.append(cleaned)

            if not text and not tables:
                empty_pages.append(page_num)

            pages.append(ExtractedPage(
                page_number=page_num,
                text=text,
                tables=tables,
                images=[],
            ))
            logger.debug(
                "  page {} | chars={} | tables={}",
                page_num, len(text), len(tables),
            )

        if empty_pages:
            logger.warning(
                "{} page(s) had no extractable content: {}",
                len(empty_pages), empty_pages[:20],
            )

        logger.info(
            "Claude PDF parsed | file={} | pages={} | total_tables={}",
            path.name, len(pages), sum(len(p.tables) for p in pages),
        )

        return ExtractedDocument(
            file_name=path.name,
            mime_type="application/pdf",
            num_pages=len(pages),
            pages=pages,
        )

    @retry(
        retry=retry_if_exception(_is_retryable_anthropic_error),
        wait=wait_exponential(multiplier=1, min=2, max=60),
        stop=stop_after_attempt(5),
        reraise=True,
    )
    def _call_claude_stream_with_retry(
        self,
        *,
        client: Any,
        messages: list[dict[str, Any]],
    ) -> Any:
        with _PDF_EXTRACTION_SEMAPHORE:
            with client.messages.stream(
                model=settings.claude_pdf_model,
                max_tokens=32768,
                messages=messages,
            ) as stream:
                return stream.get_final_message()

    # ---------- DOCX ----------
    def _extract_docx(
        self, path: Path, doc_id: str, image_dir: Path
    ) -> ExtractedDocument:
        """Extract DOCX with page awareness + inline image extraction.

        python-docx has no page concept, but DOCX files mark page breaks
        explicitly with <w:br w:type="page"/> and implicitly at section
        breaks. We walk the body element in document order, accumulate
        paragraphs, tables, and images into the current page, and start
        a new page whenever we hit a page break.

        Images: each <w:drawing> inside a run references an image part
        by rId. We resolve the rId against `doc.part.related_parts` to
        get the raw image bytes, save them, and (if OCR is available)
        extract their text. Each image becomes an `ExtractedImage` on
        the page it appears on.
        """
        from docx import Document as DocxDocument
        from docx.oxml.ns import qn
        from docx.table import Table as DocxTable
        from docx.text.paragraph import Paragraph as DocxParagraph

        d = DocxDocument(str(path))
        body = d.element.body

        # Resolve image parts by rId ONCE so we don't repeat work per image
        related_parts = d.part.related_parts

        # Accumulators for the CURRENT page being built
        current_page_num = 1
        current_paragraphs: list[str] = []
        current_tables: list[list[list[str]]] = []
        current_images: list[ExtractedImage] = []
        pages: list[ExtractedPage] = []
        global_img_counter = 0  # stable image index across the whole doc

        def flush_page() -> None:
            nonlocal current_paragraphs, current_tables, current_images, current_page_num
            pages.append(ExtractedPage(
                page_number=current_page_num,
                text="\n\n".join(current_paragraphs),
                tables=current_tables,
                images=current_images,
            ))
            current_page_num += 1
            current_paragraphs = []
            current_tables = []
            current_images = []

        def paragraph_has_page_break(p_elem) -> bool:
            for br in p_elem.iter(qn("w:br")):
                if br.get(qn("w:type")) == "page":
                    return True
            if p_elem.find(".//" + qn("w:lastRenderedPageBreak")) is not None:
                return True
            return False

        def section_has_page_break(p_elem) -> bool:
            return p_elem.find(qn("w:pPr") + "/" + qn("w:sectPr")) is not None

        def extract_images_from_paragraph(p_elem) -> list[ExtractedImage]:
            """Find all <a:blip r:embed='rId...'> inside this paragraph,
            materialize them as ExtractedImage objects, and inject any
            vision-extracted tables into the current page."""
            nonlocal global_img_counter
            found: list[ExtractedImage] = []
            blip_tag = (
                "{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
            )
            embed_attr = (
                "{http://schemas.openxmlformats.org/officeDocument/"
                "2006/relationships}embed"
            )
            for blip in p_elem.iter(blip_tag):
                rid = blip.get(embed_attr)
                if not rid or rid not in related_parts:
                    continue
                part = related_parts[rid]
                try:
                    data = part.blob
                except Exception as e:
                    logger.warning("DOCX image rId={} blob read failed: {}", rid, e)
                    continue

                global_img_counter += 1
                content_type = getattr(part, "content_type", "") or ""
                fmt = "png"
                for ext in ("png", "jpeg", "jpg", "gif", "tiff", "bmp"):
                    if ext in content_type.lower():
                        fmt = "jpg" if ext == "jpeg" else ext
                        break

                img, injected = self._process_image_bytes(
                    data=data,
                    fmt=fmt,
                    page_number=current_page_num,
                    image_index=global_img_counter,
                    doc_id=doc_id,
                    image_dir=image_dir,
                    source="embedded",
                )
                if img is not None:
                    found.append(img)
                # Any tables found inside this image get added to the
                # current page so the chunker turns them into table_row chunks.
                if injected:
                    current_tables.extend(injected)
            return found

        # Walk body children in XML order
        for child in body.iterchildren():
            tag = child.tag

            if tag == qn("w:p"):
                para = DocxParagraph(child, d)
                text = (para.text or "").strip()

                had_break_before = False
                for run in child.iter(qn("w:r")):
                    first_break = False
                    for run_child in run:
                        local = run_child.tag.split("}")[-1]
                        if local == "br" and run_child.get(qn("w:type")) == "page":
                            first_break = True
                            break
                        if local == "t" and (run_child.text or "").strip():
                            break
                    if first_break:
                        had_break_before = True
                        break

                if had_break_before and (
                    current_paragraphs or current_tables or current_images
                ):
                    flush_page()

                if text:
                    current_paragraphs.append(text)

                # Extract images anchored in this paragraph
                current_images.extend(extract_images_from_paragraph(child))

                if paragraph_has_page_break(child) and not had_break_before:
                    flush_page()
                elif section_has_page_break(child):
                    flush_page()

            elif tag == qn("w:tbl"):
                tbl = DocxTable(child, d)
                rows: list[list[str]] = []
                for row in tbl.rows:
                    cells = [(cell.text or "").strip() for cell in row.cells]
                    if any(cells):
                        rows.append(cells)
                if rows:
                    current_tables.append(rows)
                # Images may also be embedded inside table cells
                current_images.extend(extract_images_from_paragraph(child))

        if current_paragraphs or current_tables or current_images or not pages:
            flush_page()

        # Prune empty trailing pages, but keep at least one.
        while len(pages) > 1 and not pages[-1].text and not pages[-1].tables and not pages[-1].images:
            pages.pop()

        logger.info(
            "DOCX pages detected | file={} | pages={} | tables={} | images={}",
            path.name,
            len(pages),
            sum(len(p.tables) for p in pages),
            sum(len(p.images) for p in pages),
        )

        return ExtractedDocument(
            file_name=path.name,
            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            num_pages=len(pages),
            pages=pages,
        )

    # ---------- Plain text ----------
    def _extract_txt(self, path: Path) -> ExtractedDocument:
        text = path.read_text(encoding="utf-8", errors="replace")
        return ExtractedDocument(
            file_name=path.name,
            mime_type="text/plain",
            num_pages=1,
            pages=[ExtractedPage(page_number=1, text=text)],
        )

    def _extract_image(self, path: Path, document_id: str, image_dir: Path) -> ExtractedDocument:
        """Standalone image or offline scan (PNG/JPEG/TIFF/etc.) via vision + OCR fallback."""
        data = path.read_bytes()
        ext = path.suffix.lower().lstrip(".")
        vision_fmt = "jpeg" if ext in ("jpg", "jpeg") else ext
        if vision_fmt == "tif":
            vision_fmt = "tiff"

        ocr_text = ocr_service.ocr_bytes(data) if ocr_service.available else ""
        vision_text = ""
        tables: list[list[list[str]]] = []
        if vision_service.available:
            vr = vision_service.extract_from_image(data, image_format=vision_fmt)
            vision_text = vr.combined_text
            tables = [t for t in vr.tables if t]
        elif not ocr_text:
            logger.warning(
                "Image extraction has no vision or OCR | file={}", path.name
            )

        page_text = "\n\n".join(p for p in (vision_text, ocr_text) if p).strip()
        fmt = ext if ext != "jpeg" else "jpg"
        saved_path, url_path = self._save_image_bytes(
            data, image_dir, document_id, 1, 1, fmt or "png"
        )
        image = ExtractedImage(
            page_number=1,
            image_index=1,
            ocr_text=page_text,
            format=fmt,
            saved_path=saved_path or None,
            url_path=url_path or None,
            source="scan",
        )
        mime = f"image/{fmt}" if fmt else "image/png"
        return ExtractedDocument(
            file_name=path.name,
            mime_type=mime,
            num_pages=1,
            pages=[
                ExtractedPage(
                    page_number=1,
                    text=page_text,
                    tables=tables,
                    images=[image],
                )
            ],
        )


extraction_service = ExtractionService()
